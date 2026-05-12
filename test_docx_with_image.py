#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试包含图片的 Word 文档解析
"""

import sys
from pathlib import Path

# 添加项目根目录到路径
root_dir = Path(__file__).parent
sys.path.insert(0, str(root_dir))

from app.tools.file_reader import read_docx
from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont
import os


def create_simple_image():
    """创建一个简单的测试图片"""
    # 创建一个简单的图片
    img = Image.new('RGB', (400, 200), color='white')
    draw = ImageDraw.Draw(img)
    
    # 画一些简单的图形和文字
    draw.rectangle([50, 50, 350, 150], outline='blue', width=3)
    draw.text((150, 80), "测试图片", fill='black')
    
    # 保存图片
    img_path = root_dir / 'test_files' / 'test_image.png'
    img.save(img_path)
    return img_path


def create_test_docx_with_image():
    """创建一个包含图片的 Word 文档"""
    doc = Document()
    
    # 添加标题
    doc.add_heading('包含图片的测试文档', 0)
    
    # 添加段落
    doc.add_paragraph('这是一个包含图片的测试文档。')
    
    # 创建并添加图片
    img_path = create_simple_image()
    doc.add_picture(str(img_path), width=Inches(3.0))
    
    # 添加更多内容
    doc.add_paragraph()
    doc.add_paragraph('图片下方的文字内容。')
    
    # 保存文档
    test_file = root_dir / 'test_files' / 'test_with_image.docx'
    doc.save(test_file)
    print(f"[OK] 包含图片的测试文档已创建: {test_file}")
    return test_file


def test_docx_parser(file_path):
    """测试 Word 文档解析"""
    print("\n" + "="*60)
    print("测试包含图片的 Word 文档解析器")
    print("="*60)
    
    result = read_docx(str(file_path))
    
    print(f"\n[FILE] 文件解析结果:")
    print(f"   文本内容长度: {len(result['text'])} 字符")
    print(f"   表格数量: {len(result['tables'])}")
    print(f"   元数据: {result['metadata']}")
    
    print("\n[TEXT] 提取的文本:")
    print("-" * 60)
    print(result['text'][:500])
    print("-" * 60)
    
    # 检查是否检测到图片
    image_count = result['metadata'].get('image_count', 0)
    if image_count > 0:
        print(f"\n[IMAGE] 检测到 {image_count} 张图片！")
    else:
        print("\n[WARNING] 未检测到图片")
    
    print("\n[OK] 测试完成!")


if __name__ == "__main__":
    # 创建测试文档
    test_file = create_test_docx_with_image()
    
    # 测试解析
    test_docx_parser(test_file)
