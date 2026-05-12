#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试 Word 文档解析器
"""

import sys
from pathlib import Path

# 添加项目根目录到路径
root_dir = Path(__file__).parent
sys.path.insert(0, str(root_dir))

from app.tools.file_reader import read_docx
from docx import Document
from docx.shared import Inches


def create_test_docx():
    """创建一个测试用的 Word 文档"""
    doc = Document()
    
    # 添加标题
    doc.add_heading('测试文档', 0)
    
    # 添加段落
    doc.add_paragraph('这是一个包含文字的测试文档。')
    doc.add_paragraph('下面是一些内容：')
    
    # 添加一个简单的表格
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = '项目'
    table.rows[0].cells[1].text = '说明'
    table.rows[1].cells[0].text = '测试'
    table.rows[1].cells[1].text = '成功'
    
    # 添加一些空段落模拟有图片的文档
    doc.add_paragraph()
    doc.add_paragraph('[这里应该有一张图片]')
    doc.add_paragraph()
    
    # 保存文档
    test_file = root_dir / 'test_files' / 'test_document.docx'
    doc.save(test_file)
    print(f"[OK] 测试文档已创建: {test_file}")
    return test_file


def test_docx_parser(file_path):
    """测试 Word 文档解析"""
    print("\n" + "="*60)
    print("测试 Word 文档解析器")
    print("="*60)
    
    result = read_docx(str(file_path))
    
    print(f"\n[FILE] 文件解析结果:")
    print(f"   文本内容长度: {len(result['text'])} 字符")
    print(f"   表格数量: {len(result['tables'])}")
    print(f"   元数据: {result['metadata']}")
    
    print("\n[TEXT] 提取的文本:")
    print("-" * 60)
    print(result['text'][:500])  # 只打印前500字符
    print("-" * 60)
    
    if result['tables']:
        print(f"\n[TABLE] 提取的表格（共 {len(result['tables'])} 个）:")
        for i, table in enumerate(result['tables']):
            print(f"\n  表格 {i+1}:")
            for row in table:
                print(f"    {row}")
    
    print("\n[OK] 测试完成!")


if __name__ == "__main__":
    # 创建测试文档
    test_file = create_test_docx()
    
    # 测试解析
    test_docx_parser(test_file)
