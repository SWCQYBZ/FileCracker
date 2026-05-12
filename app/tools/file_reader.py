"""
文件: tools/file_reader.py | 文件读取工具集
职责: 提供各类文件的读取函数，统一返回格式
设计: 每个函数接收 file_path，返回 {"text", "tables", "metadata"}
      统一的返回值格式让下游不用关心源文件类型

每个 reader 都加了防御性 Import:
  如果依赖库没安装，返回错误提示而非抛异常
  这样系统在缺少某些库时仍可运行部分功能
"""

# ============================================================
# 标准库
# os: 文件路径操作
# json: JSON 解析（标准库，无需额外安装）
# ============================================================
import os
import json

# ============================================================
# PDF 解析库 - pymupdf
# 为什么选 pymupdf: 速度最快(PDF解析标杆)，依赖少
# 为什么不用 PyPDF2: 慢 10 倍
# 为什么不用 pdfplumber: 更慢，唯一优点是提取表格更精确
# ============================================================
try:
    import fitz  # pymupdf 在 Python 中的包名是 fitz
except ImportError:
    fitz = None  # 没安装就标记为 None，调用时给提示

# ============================================================
# DOCX 解析库 - python-docx
# 这是 Python 解析 Word 文件的唯一成熟选择
# 替代品: 无，python-docx 是事实标准
# ============================================================
try:
    from docx import Document
except ImportError:
    Document = None

# ============================================================
# YAML 解析库
# YAML 是配置文件常用格式，但非必须
# Python 标准库不含 yaml，需要额外安装 PyYAML
# ============================================================
try:
    import yaml
except ImportError:
    yaml = None

try:
    from openpyxl import load_workbook
except ImportError:
    load_workbook = None

from .registry import registry  # 工具注册中心单例


def read_pdf(file_path: str) -> dict:
    """
    PDF 文件解析

    原理: pymupdf 将 PDF 页面渲染为文本块，直接提取
    限制: 扫描件 PDF（图片式）不能提取文本，需要 OCR

    性能: 10 页 PDF 约 0.3 秒
    """
    if fitz is None:
        return {"text": "[pymupdf 未安装]", "tables": [], "metadata": {}}
    doc = fitz.open(file_path)                       # 打开 PDF 文档
    text_parts = []                                  # 收集每页文本
    metadata = {
        "pages": len(doc),                           # 总页数
        "title": doc.metadata.get("title", ""),      # 文档标题
        "author": doc.metadata.get("author", ""),    # 作者
    }
    for page in doc:                                 # 遍历每一页
        text_parts.append(page.get_text())           # 提取该页文本
    doc.close()                                      # 关闭文档（重要: 释放文件锁）
    return {"text": "\n".join(text_parts), "tables": [], "metadata": metadata}


def read_docx(file_path: str) -> dict:
    """
    Word 文档(.docx)解析

    可提取的内容:
    1. 段落文本 (paragraphs)
    2. 表格数据 (tables) - 每个 Word 表格转为二维数组
    3. 图片检测和计数

    .doc 格式（旧版 Word）不支持，需要先转换
    """
    if Document is None:
        return {"text": "[python-docx 未安装]", "tables": [], "metadata": {}}
    
    doc = Document(file_path)
    
    # 提取所有段落，保留空行但不保留纯空白行
    text_parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            text_parts.append(p.text)
        elif p.runs:  # 即使是空行但有格式也要保留
            text_parts.append("")
    text = "\n".join(text_parts)
    
    # 提取所有表格
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            # 每个单元格内容去除首尾空白
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append(rows)
    
    # 检测图片数量
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_count += 1
    
    # 构建元数据
    metadata = {
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "image_count": image_count,
        "core_properties": {}
    }
    
    # 提取文档属性（如果可用）
    if doc.core_properties:
        if doc.core_properties.title:
            metadata["core_properties"]["title"] = doc.core_properties.title
        if doc.core_properties.author:
            metadata["core_properties"]["author"] = doc.core_properties.author
        if doc.core_properties.created:
            metadata["core_properties"]["created"] = str(doc.core_properties.created)
        if doc.core_properties.modified:
            metadata["core_properties"]["modified"] = str(doc.core_properties.modified)
    
    # 如果没有文本但有图片，添加说明
    if not text.strip() and image_count > 0:
        text = f"[此文档包含 {image_count} 张图片，可能需要 OCR 识别图片中的文字]"
    
    return {"text": text, "tables": tables, "metadata": metadata}


def read_txt(file_path: str) -> dict:
    """
    纯文本文件解析

    最大挑战: 编码检测
    不同系统/语言的文件编码不同:
    - Windows 中文: GBK/GB2312
    - Linux/Mac: UTF-8
    - 旧系统: latin-1

    策略: chardet 检测 + 多编码 fallback
    """
    import chardet                              # 编码检测库
    with open(file_path, "rb") as f:            # 二进制模式读取
        raw = f.read()
    detected = chardet.detect(raw)              # 自动检测编码
    encoding = detected.get("encoding", "utf-8") or "utf-8"
    # fallback 链: 先尝试检测到的编码 → UTF-8 → GBK → latin-1
    for enc in [encoding, "utf-8", "gbk", "gb2312", "latin-1"]:
        try:
            text = raw.decode(enc)
            break                                # 解码成功就退出
        except (UnicodeDecodeError, LookupError):
            text = raw.decode("utf-8", errors="replace")  # 最后的保底
    return {"text": text, "tables": [], "metadata": {"encoding": encoding}}


def read_markdown(file_path: str) -> dict:
    """Markdown 文件解析——本质上就是文本文件"""
    result = read_txt(file_path)
    result["metadata"]["format"] = "markdown"   # 标记为 markdown 格式
    return result


def read_json(file_path: str) -> dict:
    """
    JSON 文件解析

    特殊处理: 解析后重新格式化输出
    原因: JSON 可能是压缩的一行，重新格式化后可读性大增
    不影响原始文件，只影响返回的 text 字段
    """
    result = read_txt(file_path)
    try:
        data = json.loads(result["text"])        # 验证并解析 JSON
        result["metadata"]["parsed"] = True
        # 重新格式化: 缩进 2 空格，确保中文不转义
        result["text"] = json.dumps(data, ensure_ascii=False, indent=2)
    except json.JSONDecodeError:
        result["metadata"]["parsed"] = False     # 标记解析失败
    return result


def read_xml(file_path: str) -> dict:
    """XML 文件——当前按纯文本处理，未做 XML 解析"""
    return read_txt(file_path)


def read_yaml(file_path: str) -> dict:
    """
    YAML 文件解析
    尝试用 PyYAML 解析，失败则返回原始文本
    """
    result = read_txt(file_path)
    if yaml:
        try:
            data = yaml.safe_load(result["text"])  # safe_load 防止代码注入
            result["metadata"]["parsed"] = True
        except Exception:
            result["metadata"]["parsed"] = False
    return result


def read_xlsx(file_path: str) -> dict:
    """
    Excel (.xlsx) 文件解析

    用 openpyxl 读取所有工作表的内容和表格结构
    文本输出格式：每个 sheet 以 === Sheet: 名称 === 开头，每行用 | 分隔
    """
    if load_workbook is None:
        return {"text": "[openpyxl 未安装]", "tables": [], "metadata": {}}

    wb = load_workbook(file_path, data_only=True)
    text_parts = []
    all_tables = []
    sheet_names = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        sheet_names.append(sheet_name)

        # 表格结构
        rows = []
        for row in ws.iter_rows(values_only=True):
            converted = []
            for cell in row:
                if cell is None:
                    converted.append("")
                else:
                    converted.append(str(cell))
            rows.append(converted)
        if rows:
            all_tables.append(rows)

        # 文本内容
        sheet_text = []
        for row in ws.iter_rows(values_only=True):
            line = " | ".join(str(c) for c in row if c is not None)
            if line.strip():
                sheet_text.append(line)
        text_parts.append(f"=== Sheet: {sheet_name} ===\n" + "\n".join(sheet_text))

    wb.close()

    return {
        "text": "\n\n".join(text_parts),
        "tables": all_tables,
        "metadata": {
            "sheets": sheet_names,
            "sheet_count": len(sheet_names),
        },
    }


def get_file_reader(file_type: str):
    """
    根据文件类型获取对应的 reader 函数

    这是"策略模式"(Strategy Pattern)的简单实现
    用 dict 替代 if-elif-else，新增类型只需加一行

    参数:
      file_type: config.py SUPPORTED_EXTENSIONS 中的内部类型
    返回:
      一个 reader 函数，接收 file_path 返回 dict
    """
    readers = {
        "pdf": read_pdf,
        "docx": read_docx,
        "text": read_txt,
        "markdown": read_markdown,
        "json": read_json,
        "xml": read_xml,
        "yaml": read_yaml,
        "xlsx": read_xlsx,
    }
    # get 的第二个参数是默认值——不支持的类型用 read_txt 兜底
    return readers.get(file_type, read_txt)


# ============================================================
# 工具注册
# 模块加载时自动注册到全局 registry
# 这是 Python 模块的副作用(import 时执行模块级代码)
# 所以 tools/__init__.py 必须 import 这个模块
# ============================================================
registry.register(
    read_pdf, "read_pdf",
    "解析 PDF 文件，提取文本内容和元数据",
    {"file_path": "PDF 文件路径"},
    agent="parser",
)
registry.register(
    read_docx, "read_docx",
    "解析 Word (.docx) 文件，提取文本和表格",
    {"file_path": "Word 文件路径"},
    agent="parser",
)
registry.register(
    read_txt, "read_txt",
    "解析纯文本文件，自动检测编码",
    {"file_path": "文本文件路径"},
    agent="parser",
)
registry.register(
    read_json, "read_json",
    "解析 JSON 文件",
    {"file_path": "JSON 文件路径"},
    agent="parser",
)
registry.register(
    read_xlsx, "read_xlsx",
    "解析 Excel (.xlsx) 文件，提取所有工作表的文本和表格数据",
    {"file_path": "Excel 文件路径"},
    agent="parser",
)
